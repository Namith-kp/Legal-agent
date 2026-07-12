import os
import re
import sys
import glob
from dataclasses import dataclass
import tiktoken
import chromadb
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction

# Initialize tiktoken encoding
encoding = tiktoken.get_encoding("cl100k_base")

# Database persistence directory
CHROMA_DB_DIR = os.environ.get("CHROMA_DB_DIR", os.path.join("data", "chroma_db"))

@dataclass
class Chunk:
    text: str
    metadata: dict

def get_chroma_client():
    """Returns a persistent Chroma DB client."""
    os.makedirs(CHROMA_DB_DIR, exist_ok=True)
    return chromadb.PersistentClient(path=CHROMA_DB_DIR)

def get_embedding_function():
    """Returns the default embedding function for RAG (all-MiniLM-L6-v2)."""
    return SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")

def extract_citation_metadata(text: str, default_page: int = 1, default_act: str = "Unknown Act") -> dict:
    """
    Extracts citation metadata (act name, section, page) from the text.
    Uses default fallback parameters if not explicitly found in text.
    """
    # Clean the default_act filename to act name
    clean_act = os.path.splitext(os.path.basename(default_act))[0]
    clean_act = clean_act.replace("_", " ").replace("-", " ")
    
    metadata = {
        "act": clean_act,
        "section": "Unknown Section",
        "page": str(default_page),
        "citation": ""
    }
    
    # 1. Look for common Indian legal acts/codes with their years first
    common_acts = [
        "Indian Penal Code", "Code of Criminal Procedure", "Code of Civil Procedure",
        "Indian Evidence Act", "Constitution of India", "Information Technology Act",
        "Companies Act", "Indian Contract Act", "Specific Relief Act", "Limitation Act"
    ]
    
    found_act = False
    for act in common_acts:
        # Check if the common act name appears with its year (comma or space separated)
        match = re.search(rf'\b({re.escape(act)}(?:,\s+|\s+)\d{{4}})\b', text, re.IGNORECASE)
        if match:
            metadata["act"] = match.group(1).strip()
            found_act = True
            break
            
    if not found_act:
        # Try a general regex pattern for any capitalized Act/Code/Constitution/etc. name with a year
        general_match = re.search(
            r'\b([A-Z][a-zA-Z\s]+(?:Act|Code|Constitution|Procedure|Evidence|Penal|Agreement|Contract|Law)[a-zA-Z\s]*,\s+\d{4})\b',
            text
        )
        if general_match:
            metadata["act"] = general_match.group(1).strip()
            found_act = True
            
    if not found_act:
        # Check if any common act name without a year is present
        for act in common_acts:
            if act.lower() in text.lower():
                metadata["act"] = act
                break

    # 2. Extract Section (e.g. Section 302, Sec. 378, Sec 4)
    sec_match = re.search(r'\b(?:Section|Sec\.?)\s+(\d+[A-Z]*)', text, re.IGNORECASE)
    if sec_match:
        metadata["section"] = f"Section {sec_match.group(1)}"
        
    # 3. Extract Page (e.g. Page 12, Pg 5)
    page_match = re.search(r'\b(?:Page|Pg\.?)\s+(\d+)', text, re.IGNORECASE)
    if page_match:
        metadata["page"] = page_match.group(1)
        
    # Build citation string
    parts = []
    if metadata["act"] and metadata["act"] != "Unknown Act":
        parts.append(metadata["act"])
    if metadata["section"] != "Unknown Section":
        parts.append(metadata["section"])
    if metadata["page"]:
        parts.append(f"Page {metadata['page']}")
        
    metadata["citation"] = ", ".join(parts) if parts else "Unknown Source"
    return metadata

def chunk_text_by_tokens(text: str, filename: str, default_page: int = 1) -> list[Chunk]:
    """
    Chunks standard string text into 500-token chunks with 50-token overlap.
    Extracts citation metadata for each chunk.
    """
    tokens = encoding.encode(text)
    num_tokens = len(tokens)
    chunks = []
    
    if num_tokens == 0:
        return []
        
    if num_tokens <= 500:
        chunk_text = encoding.decode(tokens)
        meta = extract_citation_metadata(chunk_text, default_page=default_page, default_act=filename)
        chunks.append(Chunk(text=chunk_text, metadata=meta))
        return chunks
        
    step = 450 # 500 (chunk size) - 50 (overlap)
    for i in range(0, num_tokens, step):
        chunk_tokens = tokens[i : i + 500]
        if len(chunk_tokens) < 10 and i > 0:
            break
        chunk_text = encoding.decode(chunk_tokens)
        meta = extract_citation_metadata(chunk_text, default_page=default_page, default_act=filename)
        chunks.append(Chunk(text=chunk_text, metadata=meta))
        
        if i + 500 >= num_tokens:
            break
            
    return chunks

def chunk_pdf(file_path: str) -> list[Chunk]:
    """Extracts text page by page from a PDF and chunks each page."""
    import pypdf
    reader = pypdf.PdfReader(file_path)
    filename = os.path.basename(file_path)
    
    chunks = []
    for page_idx, page in enumerate(reader.pages):
        page_num = page_idx + 1
        page_text = page.extract_text() or ""
        page_text = page_text.strip()
        if not page_text:
            continue
            
        page_chunks = chunk_text_by_tokens(page_text, filename, default_page=page_num)
        chunks.extend(page_chunks)
        
    return chunks

def chunk_docx(file_path: str) -> list[Chunk]:
    """Extracts paragraphs and tables from a DOCX and chunks the text."""
    import docx
    doc = docx.Document(file_path)
    filename = os.path.basename(file_path)
    
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
            
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
                
    full_text = "\n\n".join(text_parts)
    return chunk_text_by_tokens(full_text, filename, default_page=1)

def chunk_xlsx(file_path: str) -> list[Chunk]:
    """Extracts sheet rows from an XLSX and chunks the text."""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    filename = os.path.basename(file_path)
    
    text_parts = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_text_parts = []
        for row in sheet.iter_rows(values_only=True):
            row_vals = [str(val).strip() for val in row if val is not None]
            if row_vals:
                sheet_text_parts.append(" | ".join(row_vals))
        if sheet_text_parts:
            text_parts.append(f"Sheet: {sheet_name}\n" + "\n".join(sheet_text_parts))
            
    full_text = "\n\n".join(text_parts)
    return chunk_text_by_tokens(full_text, filename, default_page=1)

def chunk_document(file_path: str) -> list[Chunk]:
    """Parses and chunks PDF, DOCX, and XLSX files."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return chunk_pdf(file_path)
    elif ext == ".docx":
        return chunk_docx(file_path)
    elif ext in [".xlsx", ".xlsm"]:
        return chunk_xlsx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def ingest_law_corpus(folder_path: str):
    """Chunks and embeds all PDFs in folder_path into persistent collection 'law_corpus'."""
    if not os.path.exists(folder_path):
        raise FileNotFoundError(f"Folder not found: {folder_path}")
        
    client = get_chroma_client()
    emb_fn = get_embedding_function()
    
    collection = client.get_or_create_collection(
        name="law_corpus",
        embedding_function=emb_fn,
        metadata={"hnsw:space": "cosine"}
    )
    
    pdf_pattern = os.path.join(folder_path, "**", "*.pdf")
    pdf_files = glob.glob(pdf_pattern, recursive=True)
    
    for pdf_file in pdf_files:
        chunks = chunk_document(pdf_file)
        if not chunks:
            continue
            
        documents = [c.text for c in chunks]
        metadatas = [c.metadata for c in chunks]
        filename = os.path.basename(pdf_file)
        ids = [f"law_corpus_{filename}_{i}" for i in range(len(chunks))]
        
        collection.upsert(
            documents=documents,
            metadatas=metadatas,
            ids=ids
        )

def ingest_session_upload(file_path: str, session_id: str):
    """Chunks and embeds a session file into collection f'session_{session_id}'."""
    client = get_chroma_client()
    emb_fn = get_embedding_function()
    
    collection = client.get_or_create_collection(
        name=f"session_{session_id}",
        embedding_function=emb_fn,
        metadata={"hnsw:space": "cosine"}
    )
    
    chunks = chunk_document(file_path)
    if not chunks:
        return
        
    documents = [c.text for c in chunks]
    metadatas = [c.metadata for c in chunks]
    filename = os.path.basename(file_path)
    ids = [f"session_{session_id}_{filename}_{i}" for i in range(len(chunks))]
    
    collection.upsert(
        documents=documents,
        metadatas=metadatas,
        ids=ids
    )

def query_collection(client, collection_name: str, query: str, top_k: int) -> list[dict]:
    """Helper to query a single collection and return normalized hits."""
    try:
        collection = client.get_collection(
            name=collection_name,
            embedding_function=get_embedding_function()
        )
    except Exception:
        return []
        
    count = collection.count()
    if count == 0:
        return []
        
    n_results = min(top_k, count)
    results = collection.query(
        query_texts=[query],
        n_results=n_results
    )
    
    hits = []
    if not results or not results["documents"] or not results["documents"][0]:
        return []
        
    documents = results["documents"][0]
    metadatas = results["metadatas"][0]
    distances = results["distances"][0]
    
    for doc, meta, dist in zip(documents, metadatas, distances):
        # Chroma Cosine distance is: 1 - CosineSimilarity.
        # similarity = 1.0 - dist
        score = 1.0 - dist
        hits.append({
            "chunk_text": doc,
            "metadata": meta,
            "score": score
        })
        
    return hits

def retrieve(query: str, session_id: str | None = None, top_k: int = 5, score_floor: float = 0.55, user_id: str | None = None):
    """
    Retrieves matching document chunks from the law corpus index and optional session index.
    Confirms user permission beforehand if user_id is provided.
    """
    # 1. SQL Role permission check if user_id provided
    if user_id:
        try:
            from src.sql import check_permission
            if not check_permission(user_id, action="generate_legal_doc"):
                return {"error": "permission_denied"}
        except (ImportError, ModuleNotFoundError):
            # Dynamic path fix for local pytest executions
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.abspath(os.path.join(current_dir, "..", ".."))
            if parent_dir not in sys.path:
                sys.path.insert(0, parent_dir)
            from src.sql import check_permission
            if not check_permission(user_id, action="generate_legal_doc"):
                return {"error": "permission_denied"}

    client = get_chroma_client()
    hits = []
    
    # 2. Query permanent Indian Law Corpus
    hits += query_collection(client, "law_corpus", query, top_k)
    
    # 3. Query temporary Session uploads
    if session_id:
        hits += query_collection(client, f"session_{session_id}", query, 3)
        
    # 4. Filter by score floor
    filtered_hits = [h for h in hits if h["score"] >= score_floor]
    
    # 5. Sort by score descending
    filtered_hits.sort(key=lambda x: x["score"], reverse=True)
    
    return [
        {
            "text": h["chunk_text"],
            "source": h["metadata"].get("citation", "Unknown Source"),
            "score": h["score"]
        }
        for h in filtered_hits
    ]
